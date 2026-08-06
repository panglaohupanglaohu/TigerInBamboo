# -*- coding: utf-8 -*-
"""生成 TigerMessenger 子项目清单 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---- 全局中文字体 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell_font(cell, text, bold=False, color=None, size=10, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_font(hdr[i], htext, bold=True, color=(255,255,255), size=10, align='center')
        shade_cell(hdr[i], '2E75B6')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            mappable = (i == len(row) - 1)
            if mappable:
                ok = val.startswith('是')
                set_cell_font(cells[i], val, bold=True, color=(0,128,0) if ok else (192,0,0), align='center')
            else:
                set_cell_font(cells[i], val, align='center' if i == 0 else 'left')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t

Y = '是 ✓'
N = '否 ✗'

# ============ 封面标题 ============
title = doc.add_heading('', level=0)
trun = title.add_run('TigerMessenger 子项目清单')
trun.font.name = '微软雅黑'
trun.font.size = Pt(22)
trun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
trun.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('场景 · 器物 · 生物 · 植物 · 开发者菜单可映射资产', size=12, color=(120,120,120)).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para(f'生成日期：{datetime.date.today().isoformat()}    代码版本：fcdd944', size=9, color=(150,150,150)).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para('说明：“菜单可放置”指该资产已登记进 src/core/buildingCatalog.js 的 BUILDING_CATALOG，'
         '可在开发者菜单(🤖)→地图编辑(🗺️)中放置/拖动/复制/改朝向，布局存 localStorage。',
         size=9, color=(120,120,120))

doc.add_paragraph()

# ============ 一、场景列表 ============
add_heading('一、场景列表', 1)
add_para('正式注册场景（src/scenes/registry.js，默认同时加载）', bold=True)
make_table(
    ['场景 ID', '名称', '文件', '说明'],
    [
        ['messenger', '信使主岛', 'scenes/messengerIsland.js', '主玩法关卡：平台、土坡、湖、码头、水晶城、湖沼'],
        ['saihoji', '西芳寺·苔寺', 'scenes/saihojiGarden.js', '苔海六景：石组如岛，沿球面参道逐景展开'],
    ],
    col_widths=[2.5, 3, 5, 7]
)
add_para('')
add_para('独立入口场景（未注册进场景系统）', bold=True)
make_table(
    ['入口', '名称', '说明'],
    [
        ['planet.html', '圆球信使 / 星球传信', 'src/planet/main.js；含红方/绿方/蓝方 NPC'],
        ['aircraft.html', '莫比斯飞船演示', '独立实验页'],
        ['retro_spaceship.html', '复古飞船演示', '独立实验页'],
        ['retro_spaceship_amber.html', '复古飞船(琥珀)演示', '独立实验页'],
        ['moebius_aircraft_outline.html', '飞船线稿演示', '独立实验页'],
    ],
    col_widths=[4.5, 4, 9]
)
add_para('')
add_para('信使主岛内地标分区（messengerIsland.js 的 landmarks）', bold=True)
make_table(
    ['分区', '说明'],
    [
        ['起始营地', '海岛悬崖瀑布营地：多层海岸/荒山山洞/崖壁叠瀑/太空水环/弹琴老人/阿狸'],
        ['月牙湖', '月牙形浅水湖 + 环湖小径 + 涟漪/水花/倒影'],
        ['背侧大湖', '球冠水域，全浅可涉'],
        ['修船厂码头', '老旧港口：渔船/起重机/货柜木箱堆/栈桥'],
        ['莫比斯水晶城', '南半球晶林大都会：母皇塔/金鳞塔/能量海/花厅鸟群'],
        ['莫比斯湖沼', '深坑月夜湖沼：巨树/萤火/水下生态（经地图编辑器放置）'],
        ['厚涂苔丘', '西芳寺缘 + 湖沼缘的厚涂苔藓地形'],
        ['云环 / 赤道风暴云墙 / 大峡谷', '天空与地形大尺度景观'],
    ],
    col_widths=[4, 13]
)
add_para('')
add_para('西芳寺苔海六景分区', bold=True)
make_table(
    ['分区 ID', '名称'],
    [
        ['moss-entry', '入口苔径'], ['master-stones', '主石之庭'],
        ['dry-cascade', '枯瀑之庭'], ['moss-islands', '苔海岛群'],
        ['empty-court', '空庭'], ['return-view', '回望石组'],
    ],
    col_widths=[4, 6]
)

doc.add_page_break()

# ============ 二、器物列表 ============
add_heading('二、器物列表（建筑 / 道具 / 载具 / 结构）', 1)
make_table(
    ['器物', '工厂来源', '菜单可放置'],
    [
        ['Hard To Find 书店', 'assets/bookshop.js', Y],
        ['水墨小房', 'assets/lowPoly.js', Y],
        ['路牌', 'assets/lowPoly.js', Y],
        ['街灯', 'assets/lowPoly.js', Y],
        ['电线杆', 'assets/lowPoly.js', Y],
        ['焦墨岩 / 黑岩', 'lowPoly.js / ancient.js createBlackRock', Y],
        ['小渔船', 'assets/harbor.js', Y],
        ['港口起重机', 'assets/harbor.js', Y],
        ['货柜木箱堆', 'assets/harbor.js', Y],
        ['修船厂码头', 'assets/harbor.js buildOldHarborScene', Y],
        ['莫比斯湖沼（地形器物）', 'world/moebiusSwamp.js', Y],
        ['莫比斯塔（母皇塔 / 金鳞沿轨塔）', 'assets/moebiusTower.js', N],
        ['莫比斯航空艇（垂绳登艇·可驾驶）', 'assets/moebiusAirship.js', N],
        ['莫比斯飞船 / 飞行器编队（可进驾驶舱）', 'assets/moebiusAircraft.js', N],
        ['基督城电车（可乘坐）', 'assets/tram.js', N],
        ['气泡座舱（绕花厅巡游）', 'assets/bubblePod.js', N],
        ['木栅栏', 'assets/lowPoly.js', N],
        ['桥', 'assets/lowPoly.js', N],
        ['低多边形树', 'assets/lowPoly.js', N],
        ['云朵', 'assets/lowPoly.js', N],
        ['水晶巨构群（InstancedMesh）', 'world/moebiusCity.js', N],
        ['金黄能量海', 'world/moebiusCity.js', N],
        ['庭园石组（立/侍/卧/桥/瀑/座六式）', 'world/saihoji.js', N],
        ['石阶 / 踏石', 'world/saihoji.js / startGarden.js', N],
        ['原住民人偶', 'world/moebiusSwamp.js buildNativeDoll', N],
        ['贝壳', 'world/moebiusSwamp.js buildShell', N],
        ['莲叶舟', 'world/moebiusSwamp.js（swamp-lotus-leaf-boat）', N],
        ['池水 / 洪隐山石壁 / 叠水瀑布 / 苔岩岛', 'world/startGarden.js', N],
        ['太空水环 / 崖壁叠瀑 / 山洞', 'world/startingCamp.js', N],
    ],
    col_widths=[6, 7, 3.5]
)

doc.add_page_break()

# ============ 三、生物列表 ============
add_heading('三、生物列表（动物 / 角色 NPC）', 1)
make_table(
    ['生物 / 角色', '工厂来源', '菜单可放置'],
    [
        ['阿狸·小狐狸（可跟随 / 对话）', 'assets/fox.js', Y],
        ['赛博水墨虎（巡游 + 下坑饮水）', 'world/moebiusTiger.js', N],
        ['Boids 小鸟群（峡谷高空）', 'world/flock.js createLowPolyBird', N],
        ['花厅楼顶鸟群（母皇塔尖环绕）', 'world/flock.js（hallFlock）', N],
        ['花厅巡航鸟 / 送别伴飞编队', 'world/moebiusCity.js makeBird', N],
        ['异星滑翔长翼鸟（航空艇护航队）', 'world/airshipEscort.js', N],
        ['丹顶鹤 NPC / 岩上鹤', 'assets/ancient.js', N],
        ['弹琴老人（山洞旁 NPC）', 'world/startingCamp.js', N],
        ['送信人（玩家角色）', 'player/messenger.js', N],
        ['数字孪生送信人', 'player/agentMessenger.js', N],
        ['任务 NPC：小虎', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：阿竹', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：月见', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：星野', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：驿站', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：远方', 'quest/questSystem.js + quest/npc.js', N],
        ['任务 NPC：月影', 'quest/questSystem.js + quest/npc.js', N],
        ['planet NPC：红方 / 绿方 / 蓝方', 'planet/npcs.js', N],
        ['沼泽·白鲸（昂首破水）', 'world/moebiusSwamp.js buildBelugaWhale', N],
        ['沼泽·黄绿鳗形生物', 'world/moebiusSwamp.js buildSwampEel', N],
        ['沼泽·橙红管状蠕虫丛', 'world/moebiusSwamp.js buildTubeWormCluster', N],
        ['沼泽·粉色长尾垂挂生物（蛞蝓/水母感）', 'world/moebiusSwamp.js buildPinkHanger', N],
        ['沼泽·沼泽鸟', 'world/moebiusSwamp.js buildSwampBird', N],
        ['沼泽·长尾猴（投果互动）', 'world/moebiusSwamp.js buildLongTailMonkey', N],
        ['沼泽·发光蜥蜴', 'world/moebiusSwamp.js buildGlowLizard', N],
        ['沼泽·发光带鱼', 'world/moebiusSwamp.js buildRibbonFish', N],
        ['沼泽·绿黑斑纹小鱼群', 'world/moebiusSwamp.js（内联）', N],
        ['沼泽·萤火虫', 'world/moebiusSwamp.js（辉光 sprite）', N],
    ],
    col_widths=[7, 7.5, 2.5]
)

doc.add_page_break()

# ============ 四、植物列表 ============
add_heading('四、植物列表', 1)
make_table(
    ['植物', '工厂来源', '菜单可放置'],
    [
        ['古松（分形水墨）', 'assets/ancient.js createAncientPineTree', Y],
        ['绣球花丛', 'assets/hydrangea.js', Y],
        ['水墨小花', 'assets/lowPoly.js', Y],
        ['草坪山丘', 'assets/lowPoly.js', Y],
        ['草丛', 'assets/bookshop.js createGrassTuft', N],
        ['书店绣球（变体）', 'assets/hydrangea.js createBookshopHydrangeas', N],
        ['红枫', 'world/startGarden.js createMaple', N],
        ['竹林（竹竿 + 介字撇叶簇）', 'world/startGarden.js createBambooWallCluster', N],
        ['厚涂苔藓地被 / 苔丘', 'world/mossyGround.js', N],
        ['苔藓斑块 / 石底苔裙', 'world/saihoji.js / startGarden.js', N],
        ['沼泽·世界树', 'world/moebiusSwamp.js buildWorldTree', N],
        ['沼泽·苍天巨树', 'world/moebiusSwamp.js buildToweringTree', N],
        ['沼泽·坑缘棕榈 / 丛林树', 'world/moebiusSwamp.js buildRimPalm', N],
        ['沼泽·树冠顶棚（巨叶遮天）', 'world/moebiusSwamp.js buildCanopyCeiling', N],
        ['沼泽·紫色蘑菇 / 珊瑚状植物', 'world/moebiusSwamp.js buildMoebiusMushroom', N],
        ['沼泽·发光花蕊花', 'world/moebiusSwamp.js buildGlowFlower', N],
        ['沼泽·巨花', 'world/moebiusSwamp.js buildGiantFlower', N],
        ['沼泽·粉色花蕊尖锥', 'world/moebiusSwamp.js buildStamenSpike', N],
    ],
    col_widths=[7, 7.5, 2.5]
)

doc.add_page_break()

# ============ 五、开发者菜单已映射 / 可放置清单 ============
add_heading('五、开发者菜单已映射 / 可放置清单', 1)
add_para('地图编辑器（🤖 → 🗺️ 打开地图编辑）的调色板数据源 = src/core/buildingCatalog.js 的 '
         'BUILDING_CATALOG，共 16 项，全部可放置 / 拖动 / 复制 / 改朝向，布局存 localStorage。', bold=False)
add_para('')

add_para('器物（10 项）', bold=True, color=(0x1F, 0x3A, 0x5F))
make_table(
    ['目录 ID', '标签', '工厂', '标记色'],
    [
        ['bookshop', 'Hard To Find 书店', 'createHardToFindBookshop', '#c45a3a'],
        ['house', '水墨小房', 'createLowPolyHouse', '#8a9aaa'],
        ['signpost', '路牌', 'createLowPolySignpost', '#a63a2e'],
        ['lamp', '街灯', 'createLowPolyStreetLamp', '#5a6570'],
        ['pole', '电线杆', 'createLowPolyUtilityPole', '#3a322c'],
        ['rock', '焦墨岩', 'createLowPolyRock', '#4a4844'],
        ['fisherBoat', '小渔船', 'createFisherBoat', '#2C96B4'],
        ['harborCrane', '港口起重机', 'createHarborCrane', '#37474F'],
        ['stackedCrates', '货柜木箱堆', 'createStackedCrates', '#A2B5CD'],
        ['oldHarbor', '修船厂码头', 'buildOldHarborScene', '#8B7355'],
    ],
    col_widths=[3, 4.5, 5.5, 3]
)
add_para('')
add_para('植物（4 项）', bold=True, color=(0x1F, 0x3A, 0x5F))
make_table(
    ['目录 ID', '标签', '工厂', '标记色'],
    [
        ['pine', '古松', 'createAncientPineTree', '#2a4030'],
        ['hydrangea', '绣球花丛', 'createLowPolyHydrangeaBush', '#9ec5ff'],
        ['flower', '水墨小花', 'createLowPolyFlower', '#c4a090'],
        ['lawnHill', '草坪山丘', 'createLowPolyLawnHill', '#55875f'],
    ],
    col_widths=[3, 4.5, 5.5, 3]
)
add_para('')
add_para('生物（1 项）', bold=True, color=(0x1F, 0x3A, 0x5F))
make_table(
    ['目录 ID', '标签', '工厂', '标记色'],
    [['fox', '阿狸（小狐狸）', 'createLowPolyFox', '#E96A36']],
    col_widths=[3, 4.5, 5.5, 3]
)
add_para('')
add_para('地形器物（1 项）', bold=True, color=(0x1F, 0x3A, 0x5F))
make_table(
    ['目录 ID', '标签', '工厂', '标记色'],
    [['moebiusSwamp', '莫比斯湖沼', 'createMoebiusSwampPlacement', '#48C9B0']],
    col_widths=[3, 4.5, 5.5, 3]
)

add_para('')
add_heading('总结与扩展建议', 2)
add_para('开发者菜单里“能映射”（可放置）的共 16 项：器物 10、植物 4、生物 1、地形器物 1。')
add_para('其余代码已实现但未登记进 BUILDING_CATALOG、无法经地图编辑器放置的内容包括：')
for item in [
    '生物：赛博水墨虎、各类鸟群（峡谷/花厅/护航长翼鸟/沼泽鸟）、丹顶鹤、弹琴老人、送信人、7 个任务 NPC、planet 三色 NPC、沼泽 10 种生态生物',
    '植物：草丛、红枫、竹林、苔藓地被、沼泽 8 种植物（世界树/巨树/棕榈/树冠/紫蘑菇/发光花/巨花/花蕊尖锥）',
    '器物：莫比斯塔、航空艇、飞船编队、电车、气泡座舱、栅栏、桥、水晶巨构、庭园石组、人偶、贝壳、莲叶舟等',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
add_para('')
add_para('扩展方法：在 buildingCatalog.js 中新增一个 BuildingDef（引用对应工厂 + 默认朝向 / 碰撞半径 / 标记色），'
         '地图编辑器会自动把它加入调色板，即可放置。', bold=True, color=(0xC0, 0x50, 0x00))

out = r'D:\TigerInBamnoo\docs\TigerMessenger清单.docx'
doc.save(out)
print('SAVED:', out)
